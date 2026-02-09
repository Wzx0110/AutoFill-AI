import os
import shutil
import tempfile
from typing import List, Dict, Any
from fastapi import UploadFile
from pdfrw import PdfReader, PdfWriter, PdfDict, PdfName, PdfString
from docx import Document
from app.schemas.extraction import FieldResult
import logging

logger = logging.getLogger(__name__)

class FileFillerService:
    
    def fill_document(self, file: UploadFile, data: List[FieldResult]) -> str:
        """
        判斷檔案類型並執行填寫，回傳填寫後的暫存檔路徑
        """
        # 1. 轉換資料格式: List[FieldResult] -> Dict[key, value]
        data_map = {item.key: str(item.value) for item in data if item.value is not None}
        
        filename = file.filename.lower()
        
        # 建立暫存檔來操作
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp:
            shutil.copyfileobj(file.file, tmp)
            input_path = tmp.name
        
        output_path = input_path.replace(".", "_filled.")

        try:
            if filename.endswith(".pdf"):
                self._fill_pdf(input_path, output_path, data_map)
            elif filename.endswith(".docx"):
                self._fill_word(input_path, output_path, data_map)
            else:
                raise ValueError("Unsupported file format")
            
            return output_path
        finally:
            # 清理 input，保留 output 讓 controller 回傳
            if os.path.exists(input_path):
                os.remove(input_path)

    def _fill_pdf(self, input_path: str, output_path: str, data: Dict[str, str]):
        """
        針對有 AcroForm 欄位的 PDF 進行填寫
        注意: PDF 欄位名稱必須與 Extraction Field Key 一致
        """
        template_pdf = PdfReader(input_path)
        
        if template_pdf.Root.AcroForm:
            for field in template_pdf.Root.AcroForm.Fields:
                # 取得 PDF 內部的欄位名稱
                field_name = field.T
                if field_name:
                    # 處理編碼 (有些 PDF 欄位是括號包起來的)
                    key = field_name.replace("(", "").replace(")", "")
                    
                    if key in data:
                        # 填入數值
                        field.V = PdfString.encode(data[key])
                        # field.Ff = PdfObject(1) 
        
        PdfWriter().write(output_path, template_pdf)
        logger.info(f"PDF filled: {output_path}")

    def _fill_word(self, input_path: str, output_path: str, data: Dict[str, str]):
        """
        針對 Word 進行簡單的「關鍵字替換」
        它會尋找文件中的 {{key}} 並替換成 value
        """
        doc = Document(input_path)
        
        # 1. 替換段落中的文字
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}" # 尋找 {{key}}
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # 2. 替換表格中的文字
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
                            
        doc.save(output_path)
        logger.info(f"Word filled: {output_path}")

file_filler_service = FileFillerService()